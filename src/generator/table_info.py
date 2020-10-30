# -*- coding: utf-8 -*-
import docx
from src.constant.index_type import IndexType
from src.constant.constant import COL_NAMES, SCHEMA_TABLES_SQL, SCHEMA_COLS_SQL, TEMPLATE_DOCX_PATH

_author_ = 'luwt'
_date_ = '2018/12/17 13:42'


class TableInfo:

    def __init__(self,
                 conn_name,
                 table_schema,
                 table_cols,
                 path,
                 db_executor,
                 consumer,
                 template_file=TEMPLATE_DOCX_PATH):
        self.conn_name = conn_name
        self.schema = table_schema
        self.table_cols = table_cols
        self.path = path
        self.db_executor = db_executor
        self.consumer = consumer
        self.doc = docx.Document(docx=template_file)

    def get_data(self, db, sql):
        self.db_executor.switch_db(db)
        return self.db_executor.get_data(sql)

    def generate_table_title(self, table_comment, table_name):
        """生成表格标题，标题为表注释 + 表名，标题样式为有序数字"""
        self.doc.add_paragraph('\n')
        self.doc.add_paragraph(f'{table_comment} {table_name}', style='List Number')

    def generate_table(self, data):
        # Table Grid可以实现表格线框为实线
        table = self.doc.add_table(1, 6, style='Table Grid')
        # 建立表格，生成表头
        cells = table.rows[0].cells
        for i, col_name in enumerate(COL_NAMES):
            cells[i].text = col_name
        # data为查出的每个数据库表字段信息，一个字段为一个元祖
        for column_info in data:
            row_cells = table.add_row().cells
            # 表格值
            values = list(column_info)
            # 处理索引
            index = column_info[3]
            # 找对应的中文
            if index:
                values[3] = eval(f'IndexType.{index}.value')
            # 将当前字段的各个属性填充相应表格
            for i, v in enumerate(row_cells):
                if values[i]:
                    v.text = values[i]

    def main(self):
        # 查表名加表注释
        table_names_sql = f'{SCHEMA_TABLES_SQL} where table_schema = "{self.schema}";'
        # 查出所有表名与注释信息
        table_infos = self.get_data(self.schema, table_names_sql)
        table_dict = dict()
        for table_info in table_infos:
            table_dict[table_info[0]] = table_info[1]
        for table_col in self.table_cols:
            sql = SCHEMA_COLS_SQL + f"where table_schema = '{self.schema}'"
            # 整表
            if isinstance(table_col, str):
                tb_name = table_col
                sql += f" and table_name = '{tb_name}'"
            else:
                tb_name = tuple(table_col.keys())[0]
                sql += f" and table_name = '{tb_name}' and column_name in {tuple(table_col.get(tb_name))}"
            data = self.db_executor.get_data(sql)
            self.generate_table_title(table_dict.get(tb_name), tb_name)
            self.generate_table(data)
            self.consumer.send(f'[连接：{self.conn_name}] [库：{self.schema}] [表：{tb_name}]')
        self.doc.save(self.path)
