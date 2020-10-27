# -*- coding: utf-8 -*-
from get_cursor import Cursor
import index_type
import docx
import os
import time
_author_ = 'luwt'
_date_ = '2018/12/17 13:42'


# 从系统库中查出所有数据库的名称，校验输入数据库的正确性
SYS_SCHEMA_DB = "information_schema"
SCHEMA_SQL = "SELECT SCHEMA_NAME FROM `SCHEMATA`;"
# 从系统库查询tables表
SCHEMA_TABLES_SQL = "SELECT table_name, table_comment from information_schema.`TABLES`"
# 表头列名信息
col_names = ['字段名', '数据类型', '允许为空', '是否为主键', '默认值', '注释']


class TableInfo:

    def __init__(
            self, template_file, table_schema, path, host, user, pwd, port=3306, charset='utf8'
    ):
        default_docx = template_file
        self.doc = docx.Document(docx=default_docx)
        self.schema = table_schema
        self.path = path
        self.host = host
        self.user = user
        self.pwd = pwd
        self.port = port
        self.charset = charset

    def get_data(self, db, sql):
        with Cursor(
                self.host,
                self.user,
                self.pwd,
                db,
                self.port,
                self.charset
        ) as cursor:
            cursor.execute(sql)
        return cursor.fetchall()

    def generate_table_title(self, table_comment, table_name):
        """生成表格标题，标题为表注释 + 表名，标题样式为有序数字"""
        self.doc.add_paragraph('\n')
        self.doc.add_paragraph(f'{table_comment} {table_name}', style='List Number')

    def generate_table(self, data):
        # Table Grid可以实现表格线框为实线
        table = self.doc.add_table(1, 6, style='Table Grid')
        # 建立表格，生成表头
        cells = table.rows[0].cells
        for i, col_name in enumerate(col_names):
            cells[i].text = col_name
        # data为查出的每个数据库表字段信息，一个字段为一个元祖
        for column_info in data:
            row_cells = table.add_row().cells
            # 处理索引
            index = column_info[4]
            # 找对应的中文
            if column_info[4]:
                index = eval(f'index_type.IndexType.{column_info[4]}.value')
            # 表格值
            values = (
                column_info[0],
                column_info[1],
                column_info[3],
                index,
                column_info[5],
                column_info[8]
            )
            # 将当前字段的各个属性填充相应表格
            for i, v in enumerate(row_cells):
                if values[i]:
                    v.text = values[i]

    def main(self):
        # 查表名加表注释
        table_names_sql = f'{SCHEMA_TABLES_SQL} where table_schema = "{self.schema}";'
        # 查出所有表名
        table_names = self.get_data(self.schema, table_names_sql)
        for table_name in table_names:
            table_name, table_comment = table_name[0], \
                                        table_name[1].split('，')[0]
            # 查出表对象信息
            table_infos_sql = f'show full fields from {table_name};'
            data = self.get_data(self.schema, table_infos_sql)
            self.generate_table_title(table_comment, table_name)
            self.generate_table(data)
        self.doc.save(self.path)
        print('执行完毕，输出为{}，三秒后退出！'.format(self.path))
        time.sleep(3)

    def get_schema(self, schema):
        """校验数据库名输入是否正确"""
        data = self.get_data(SYS_SCHEMA_DB, SCHEMA_SQL)
        data = list(map(lambda x: str(x[0]), data))
        print(f'正确的数据库名为：{data}')
        return schema not in data


if __name__ == '__main__':
    mysql_host = 'centos121'
    mysql_user = 'root'
    mysql_pwd = 'admin'
    database = 'test'
    # 必须提供空白docx文档，作为生成文件的源文件
    template_file_path = os.path.join(os.path.abspath('.'), 'default.docx')
    out_path = 'D:\\数据词典.docx'
    table_info = TableInfo(template_file_path, database, out_path, mysql_host, mysql_user, mysql_pwd)
    if table_info.get_schema(database):
        exit("请检查数据库名，并重试")
    table_info.main()

